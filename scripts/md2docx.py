#!/usr/bin/env python3
"""Regenerate the ControlIQ PRD .docx from the authoritative Markdown source,
re-encoding the colour/shading semantics the document's annotation legend
relies on (the reverse of docx2md.py).

Usage, from the repository root:

    python3 scripts/md2docx.py \
        docs/requirement-specification/PRD.md \
        docs/requirement-specification/PRD.docx

Since v5.0 the Markdown is the authoritative source and the .docx is the
derived artefact, regenerated before each signature round. Colour mapping:

  trailing `[tag]` chips  -> coloured italic runs (legend colours below)
  ***text*** before a chip -> takes that chip's colour
  blockquote callouts      -> single-cell shaded tables (green/amber/grey/...)
  |---| tables             -> tables, header row dark blue with white text
  **[FR-xx]** feature IDs  -> blue bold

Requires python-docx.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

# Tag chip -> run colour (legend on the title page).
TAG_COLOUR = {
    '25 Jun FRD call': '1155CC',        # blue
    "Sosinna's Drive comment": 'C55A11',  # orange
    'Aug 2026 client MOM': '7030A0',    # purple
    '11 Aug 2026 MOM': '008080',        # teal
    'Aug 2026 onboarding flow': '843C0C',  # brown
    '14 Aug 2026 MOM': 'CC00CC',        # magenta
    'status note': '595959',            # grey
}
# Callout kind -> cell fill.
CALLOUT_FILL = {
    'CONFIRMED': 'E2EFDA',
    'OPEN QUESTION': 'FFF2CC',
    'DEVELOPER NOTE': 'EDEDED',
}
NOTE_FILL = {  # **NOTE** *(colour ...)* variants
    'blue': 'DDEBF7', 'orange': 'FCE4D6', 'purple': 'E6E0F5',
    'brown': 'F5E6D9', 'magenta': 'F9E0F9', 'grey': 'EDEDED',
}
HEADER_FILL = '1F4E79'
ID_COLOUR = '1155CC'
ID_RE = re.compile(r'^\[[A-Z]{2,4}-\d+[a-z]?\]$')

INLINE = re.compile(
    r'(`[^`]+`|\*\*\*.+?\*\*\*|\*\*.+?\*\*|~~.+?~~|\*[^*]+?\*|<br\s*/?>)'
)


def shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>'))


def tokenize(text, bold=False, italic=False, strike=False):
    """Yield (text, bold, italic, strike, kind) where kind is
    'chip' | 'code' | 'break' | None."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            m = re.match(r'^\[(.+)\]$', inner)
            if m and m.group(1) in TAG_COLOUR:
                yield (' [' + m.group(1) + ']', bold, True, strike, 'chip')
            else:
                yield (inner, bold, italic, strike, 'code')
        elif part.startswith('***') and part.endswith('***') and len(part) > 6:
            yield from tokenize(part[3:-3], True, True, strike)
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            yield from tokenize(part[2:-2], True, italic, strike)
        elif part.startswith('~~') and part.endswith('~~') and len(part) > 4:
            yield from tokenize(part[2:-2], bold, italic, True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            yield from tokenize(part[1:-1], bold, True, strike)
        elif re.match(r'^<br\s*/?>$', part):
            yield ('', bold, italic, strike, 'break')
        else:
            yield (part, bold, italic, strike, None)


def emit(par, text):
    """Render markdown inline text into paragraph runs, colouring italic
    segments with the colour of the tag chip that closes them."""
    toks = list(tokenize(text.replace(r'\|', '|')))
    # Assign each token the colour of the next chip in the paragraph.
    colours = [None] * len(toks)
    nxt = None
    for i in range(len(toks) - 1, -1, -1):
        if toks[i][4] == 'chip':
            nxt = TAG_COLOUR[toks[i][0].strip()[1:-1]]
            colours[i] = nxt
        else:
            colours[i] = nxt
    for (txt, bold, italic, strike, kind), colour in zip(toks, colours):
        if kind == 'break':
            par.add_run().add_break(WD_BREAK.LINE)
            continue
        run = par.add_run(txt)
        run.bold = bold
        run.italic = italic
        run.font.strike = strike
        if kind == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        if kind == 'chip':
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(colour)
        elif bold and not italic and ID_RE.match(txt):
            run.font.color.rgb = RGBColor.from_string(ID_COLOUR)
        elif italic and colour:
            run.font.color.rgb = RGBColor.from_string(colour)


def add_table(doc, rows):
    cells = [
        [c.strip() for c in re.split(r'(?<!\\)\|', r)[1:-1]]
        for r in rows
    ]
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = 'Table Grid'
    single = ncols == 1
    for i, row in enumerate(cells):
        for j in range(ncols):
            text = row[j] if j < len(row) else ''
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].text = ''
            emit(cell.paragraphs[0], text)
            if single:  # agreed / open-question summary boxes
                if text.startswith('**✅'):
                    shade(cell, CALLOUT_FILL['CONFIRMED'])
                elif text.startswith('**💬'):
                    shade(cell, CALLOUT_FILL['OPEN QUESTION'])
            elif i == 0:
                shade(cell, HEADER_FILL)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string('FFFFFF')


def callout_fill(lines):
    head = ' '.join(lines[:2])
    for kind, fill in CALLOUT_FILL.items():
        if f'**{kind}**' in head:
            return fill
    m = re.search(r'\*\*NOTE\*\*\s*\*\((\w+)', head)
    if m:
        return NOTE_FILL.get(m.group(1), 'DDEBF7')
    return 'DDEBF7'


def add_callout(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade(cell, callout_fill(lines))
    first = True
    for line in lines:
        if not line:
            continue
        par = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        if line.startswith('- '):
            par.style = doc.styles['List Bullet']
            emit(par, line[2:])
        else:
            emit(par, line)


def convert(md_path, docx_path):
    lines = open(md_path, encoding='utf-8').read().splitlines()
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    i = 0
    if lines and lines[0].startswith('<!--'):  # header comment: md-only
        while i < len(lines) and '-->' not in lines[i]:
            i += 1
        i += 1

    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            doc.add_heading(line.lstrip('#').strip(), level=min(level, 3))
            i += 1
        elif line.startswith('>'):
            block = []
            while i < len(lines) and lines[i].startswith('>'):
                block.append(lines[i].lstrip('>').strip())
                i += 1
            add_callout(doc, block)
        elif line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                if not re.match(r'^\|[\s:|-]+\|$', lines[i]):
                    rows.append(lines[i])
                i += 1
            add_table(doc, rows)
        elif line.startswith('- '):
            emit(doc.add_paragraph(style='List Bullet'), line[2:].strip())
            i += 1
        else:
            emit(doc.add_paragraph(), line.strip())
            i += 1

    doc.save(docx_path)
    print(f'wrote {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    convert(sys.argv[1], sys.argv[2])
